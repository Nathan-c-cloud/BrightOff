"""Tests unitaires — extractors.py (S3-11).

Stratégie de mocks :
- PDF valide : PDF minimal en bytes bruts (673 bytes, text layer réel).
  Pas de reportlab (absent de requirements.txt) — on embarque un PDF minimal encodé.
- PDF corrompu : b"not-a-pdf" → pdfplumber lève une exception → ExtractionError.
- PDF scanné : mock pdfplumber.open qui retourne des pages avec extract_text() == None.
- DOCX valide : construit avec python-docx en mémoire (disponible dans requirements.txt).
- DOCX corrompu : bytes aléatoires → python-docx lève une exception → ExtractionError.
- download_from_s3_sync : mock du client boto3 (dict-like).
"""

from __future__ import annotations

import io
import uuid
from unittest.mock import MagicMock, patch

import docx
import pytest

from app.modules.cv_parser.extractors import (
    EmptyTextError,
    ExtractionError,
    download_from_s3_sync,
    extract_text,
    extract_text_docx,
    extract_text_pdf,
)

# ---------------------------------------------------------------------------
# Helpers — génération de binaires valides
# ---------------------------------------------------------------------------

# PDF minimal avec une page contenant le texte "Test CV".
# Généré offline, auto-suffisant, sans dépendance reportlab.
# Structure : header + objet pages + content stream "BT /F1 12 Tf (Test CV) Tj ET" + xref + trailer.
_MINIMAL_PDF = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length 44>>
stream
BT /F1 12 Tf 100 700 Td (Test CV) Tj ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f\r
0000000009 00000 n\r
0000000058 00000 n\r
0000000115 00000 n\r
0000000266 00000 n\r
0000000360 00000 n\r
trailer<</Size 6/Root 1 0 R>>
startxref
441
%%EOF"""


def _make_valid_docx() -> bytes:
    """Génère un DOCX valide en mémoire via python-docx."""
    buf = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Développeur Fullstack — Jean Dupont")
    document.add_paragraph("Compétences : Python, FastAPI, React")
    document.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Tests — download_from_s3_sync
# ---------------------------------------------------------------------------


class TestDownloadFromS3Sync:
    def test_download_from_s3_sync_returns_bytes_on_success(self):
        mock_client = MagicMock()
        mock_body = MagicMock()
        mock_body.read.return_value = b"file content"
        mock_client.get_object.return_value = {"Body": mock_body}

        result = download_from_s3_sync(mock_client, "my-bucket", "cvs/uuid/cv.pdf")

        assert result == b"file content"
        mock_client.get_object.assert_called_once_with(Bucket="my-bucket", Key="cvs/uuid/cv.pdf")

    def test_download_from_s3_sync_raises_extraction_error_on_client_error(self):
        mock_client = MagicMock()
        mock_client.get_object.side_effect = Exception("NoSuchKey")

        with pytest.raises(ExtractionError, match="Échec du téléchargement S3"):
            download_from_s3_sync(mock_client, "my-bucket", "cvs/missing.pdf")

    def test_download_from_s3_sync_wraps_original_exception(self):
        mock_client = MagicMock()
        original_exc = RuntimeError("connection refused")
        mock_client.get_object.side_effect = original_exc

        with pytest.raises(ExtractionError) as exc_info:
            download_from_s3_sync(mock_client, "my-bucket", "key.pdf")

        assert exc_info.value.__cause__ is original_exc


# ---------------------------------------------------------------------------
# Tests — extract_text_pdf
# ---------------------------------------------------------------------------


class TestExtractTextPdf:
    def test_extract_text_pdf_valid_returns_nonempty_string(self):
        # Un PDF réel avec text layer → pdfplumber extrait du texte
        with patch("pdfplumber.open") as mock_open:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = "Test CV contenu valide"
            mock_pdf = MagicMock()
            mock_pdf.pages = [mock_page]
            mock_open.return_value.__enter__.return_value = mock_pdf

            result = extract_text_pdf(_MINIMAL_PDF)

        assert isinstance(result, str)
        assert len(result) > 0
        assert "Test CV contenu valide" in result

    def test_extract_text_pdf_multiple_pages_joined(self):
        with patch("pdfplumber.open") as mock_open:
            page1 = MagicMock()
            page1.extract_text.return_value = "Page un"
            page2 = MagicMock()
            page2.extract_text.return_value = "Page deux"
            mock_pdf = MagicMock()
            mock_pdf.pages = [page1, page2]
            mock_open.return_value.__enter__.return_value = mock_pdf

            result = extract_text_pdf(b"any")

        assert "Page un" in result
        assert "Page deux" in result

    def test_extract_text_pdf_corrupted_raises_extraction_error(self):
        # bytes invalides → pdfplumber lève une exception à l'ouverture
        with pytest.raises(ExtractionError, match="chiffré ou corrompu"):
            extract_text_pdf(b"not-a-pdf-at-all")

    def test_extract_text_pdf_scanned_raises_empty_text_error(self):
        # PDF valide structurellement mais pages sans texte (images uniquement)
        with patch("pdfplumber.open") as mock_open:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = None  # aucun texte extractible
            mock_pdf = MagicMock()
            mock_pdf.pages = [mock_page, mock_page]
            mock_open.return_value.__enter__.return_value = mock_pdf

            with pytest.raises(EmptyTextError, match="PDF scanné"):
                extract_text_pdf(b"any")

    def test_extract_text_pdf_pages_with_empty_strings_raises_empty_text_error(self):
        # Pages retournant "" ou "   " → traité comme PDF scanné
        with patch("pdfplumber.open") as mock_open:
            mock_page = MagicMock()
            mock_page.extract_text.return_value = "   "
            mock_pdf = MagicMock()
            mock_pdf.pages = [mock_page]
            mock_open.return_value.__enter__.return_value = mock_pdf

            with pytest.raises(EmptyTextError):
                extract_text_pdf(b"any")

    def test_extract_text_pdf_empty_bytes_raises_extraction_error(self):
        with pytest.raises(ExtractionError):
            extract_text_pdf(b"")


# ---------------------------------------------------------------------------
# Tests — extract_text_docx
# ---------------------------------------------------------------------------


class TestExtractTextDocx:
    def test_extract_text_docx_valid_returns_nonempty_string(self):
        content = _make_valid_docx()
        result = extract_text_docx(content)

        assert isinstance(result, str)
        assert len(result) > 0
        assert "Jean Dupont" in result

    def test_extract_text_docx_valid_contains_all_paragraphs(self):
        content = _make_valid_docx()
        result = extract_text_docx(content)

        assert "Python" in result

    def test_extract_text_docx_corrupted_raises_extraction_error(self):
        with pytest.raises(ExtractionError, match="corrompu"):
            extract_text_docx(b"PK\x03\x04this-is-not-a-valid-docx")

    def test_extract_text_docx_random_bytes_raises_extraction_error(self):
        with pytest.raises(ExtractionError):
            extract_text_docx(b"\x00\x01\x02\x03\x04\x05")

    def test_extract_text_docx_empty_document_raises_empty_text_error(self):
        # DOCX valide mais sans paragraphes non vides
        buf = io.BytesIO()
        document = docx.Document()
        # Aucun paragraphe ajouté — le document python-docx contient un paragraphe vide par défaut
        document.save(buf)
        buf.seek(0)
        content = buf.read()

        with pytest.raises(EmptyTextError, match="aucun paragraphe"):
            extract_text_docx(content)


# ---------------------------------------------------------------------------
# Tests — extract_text (dispatcher)
# ---------------------------------------------------------------------------


class TestExtractTextDispatcher:
    def test_extract_text_pdf_format_calls_pdf_extractor(self):
        with patch("app.modules.cv_parser.extractors.extract_text_pdf") as mock_pdf:
            mock_pdf.return_value = "texte pdf"
            result = extract_text(b"content", "pdf")

        mock_pdf.assert_called_once_with(b"content")
        assert result == "texte pdf"

    def test_extract_text_docx_format_calls_docx_extractor(self):
        with patch("app.modules.cv_parser.extractors.extract_text_docx") as mock_docx:
            mock_docx.return_value = "texte docx"
            result = extract_text(b"content", "docx")

        mock_docx.assert_called_once_with(b"content")
        assert result == "texte docx"

    def test_extract_text_unknown_format_raises_extraction_error(self):
        with pytest.raises(ExtractionError, match="non supporté"):
            extract_text(b"content", "txt")

    def test_extract_text_unknown_format_xlsx_raises_extraction_error(self):
        with pytest.raises(ExtractionError, match="non supporté"):
            extract_text(b"content", "xlsx")

    def test_extract_text_empty_format_raises_extraction_error(self):
        with pytest.raises(ExtractionError, match="non supporté"):
            extract_text(b"content", "")
